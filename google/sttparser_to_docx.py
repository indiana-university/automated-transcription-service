# # Speech-to-text
# Exploration and testing of the speech-to-text API

import json
import time
import argparse
import os
from google.cloud import storage
from urllib.parse import urlparse
from docx.shared import Cm, Mm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx import Document

# Common formats and styles
TABLE_STYLE_STANDARD = "Light List"

confidence_env = int(os.environ.get('CONFIDENCE', 90))

# Format output of timestamps
def timestamp(seconds):
    if (seconds.find(".") == -1):
        seconds = int(seconds[:-1])
        ts = time.gmtime(seconds)
        return time.strftime("%H:%M:%S",ts)
    else:
        x = seconds.split(".")
        seconds = int(x[0])
        ts = time.gmtime(seconds)
        return time.strftime("%H:%M:%S",ts) + "." + x[1][:3]

def decode_gcs_url(url):
    p = urlparse(url)
    return p.netloc, p.path[1:]

def download_blob(url):
    if url:
        bucket, file_path = decode_gcs_url(url)
        storage_client = storage.Client()
        bucket = storage_client.bucket(bucket)
        blob = bucket.blob(file_path)
        contents = blob.download_as_text()
        return contents

# Next version parses the final transcription result in the response file, which includes speaker ID and timestamp for every word.
def print_transcript(response, document, speakers=False):
    write_custom_text_header(document, "Audio Transcription")

    if speakers:
        last = len(response['results']) - 1
        transcript = response['results'][last]

        #Grab first first speaker and start time for initial timestamp
        best_alternative = transcript['alternatives'][0]
        try:
            current_speaker = best_alternative['words'][0]['speakerTag']
        except:
            document.add_paragraph("Speaker diarization not enabled.")
            return
        current_ts = best_alternative['words'][0]['startTime']

        document.add_paragraph()  # Spacing
        write_small_header_text(document, "WORD CONFIDENCE: >= " + str(confidence_env) + "% in black, ", (confidence_env / 100))
        write_small_header_text(document, "< " + str(confidence_env) + "% in yellow highlight", ((confidence_env - 1) / 100))

        paragraph = document.add_paragraph()
        run = paragraph.add_run()

        #Initial speaker
        run.add_text(f"[{timestamp(current_ts)}] Speaker {current_speaker}: ")
        run = paragraph.add_run()

        #Loop through words
        #When speaker changes: print line, timestamp, speaker, paragraph, add line break
        for word in best_alternative['words']:
            next_speaker = word['speakerTag']
            next_word = word['word']

            if next_speaker != current_speaker:
                #New speaker
                current_ts = word['startTime']
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_text(f"[{timestamp(current_ts)}] Speaker {current_speaker}: ")
                run = paragraph.add_run()
                current_speaker = next_speaker

            #Confidence highlighting and add word
            confidence = word['confidence']
            set_transcript_text_style(run, False, confidence=confidence)

            run.add_text(next_word + " ")
            run = paragraph.add_run()
    else:
        ts = "00:00:00"
        for result in response['results']:
            best_alternative = result['alternatives'][0]
            transcript = best_alternative.get('transcript', 'missing')
            if transcript == 'missing':
                continue
            confidence = best_alternative['confidence']
            document.add_paragraph(f"[{ts} {confidence:.0%}]: {transcript}")
            ts = timestamp(result['resultEndTime'])

def set_transcript_text_style(run, force_highlight, confidence=0.0, rgb_color=None):
    """
    Sets the colour and potentially the style of a given run of text in a transcript.  You can either
    supply the hex-code, or base it upon the confidence score in the transcript.

    :param run: DOCX paragraph run to be modified
    :param force_highlight: Indicates that we're going to forcibly set the background colour
    :param confidence: Confidence score for this word, used to dynamically set the colour
    :param rgb_color: Specific colour for the text
    """

    # If we have an RGB colour then use it
    if rgb_color is not None:
        run.font.color.rgb = rgb_color
    else:
        # Set the colour based upon the supplied confidence score
        if confidence >= (confidence_env / 100):
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Apply any other styles wanted
    if confidence == 0.0:
        # Call out any total disasters in bold
        run.font.bold = True

    # Force the background colour if required
    if force_highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def write_small_header_text(document, text, confidence):
    """
    Helper function to write out small header entries, where the text colour matches the
    colour of the transcript text for a given confidence value

    :param document: Document to write the text to
    :param text: Text to be output
    :param confidence: Confidence score, which changes the text colour
    """
    run = document.paragraphs[-1].add_run(text)
    set_transcript_text_style(run, False, confidence=confidence)
    run.font.size = Pt(10)
    run.font.italic = True

def write_custom_text_header(document, text_label):
    """
    Adds a run of text to the document with the given text label, but using our customer text-header style

    :param document: Document to write the text to
    :param text_label: Header text to write out
    :return:
    """
    paragraph = document.add_heading(text_label, level=3)

def main():
    parser = argparse.ArgumentParser('Print formatted transcipts from a speech-to-text JSON response.')
    parser.add_argument('file', metavar='file', type=str, help='a JSON file to be parsed')
    parser.add_argument('-s', '--speakers', action='store_true', dest='speakers', help='Enable speaker diarization')
    parser.add_argument('-o', '--outputFile', metavar='outputFile', type=str, help='Output DOCX file')
    args = parser.parse_args()

    document = Document()
    document.sections[0].left_margin = Mm(19.1)
    document.sections[0].right_margin = Mm(19.1)
    document.sections[0].top_margin = Mm(19.1)
    document.sections[0].bottom_margin = Mm(19.1)
    document.sections[0].page_width = Mm(210)
    document.sections[0].page_height = Mm(297)

    # Open response file and load JSON results
    if urlparse(args.file).scheme == 'gs':
        response_file = download_blob(args.file)
        response = json.loads(response_file)
    else:
        with open(args.file) as f:
            response = json.load(f)

    # Intro header
    write_custom_text_header(document, "Indiana University Social Science Research Commons")

    write_custom_text_header(document, "Google Cloud Transcribe Audio Source")
    table = document.add_table(rows=1, cols=2)
    table.style = document.styles[TABLE_STYLE_STANDARD]
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(2.0)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Job Name"
    hdr_cells[1].text = os.path.basename(__file__)
    job_data = []
    if 'results' in response and len(response['results']) > 0:
        last_result = response['results'][-2]
        dur_text = last_result.get('resultEndTime', "00:00:00")
        job_data.append({"name": "Audio Duration", "value": dur_text})
    if args.speakers:
        job_data.append({"name": "Audio Identification", "value": "Speaker-separated"})
    else:
        job_data.append({"name": "Audio Identification", "value": "Channel-separated"})
    if 'results' in response and len(response['results']) > 0:
        languages = set()
        for result in response['results']:
            languages.add(result.get('languageCode', ""))
        job_data.append({"name": "Language(s)", "value": ', '.join(filter(None, languages))})

        total_confidence = 0
        confidence_count = 0
        for result in response.get('results', []):
            for alternative in result.get('alternatives', []):
                for word in alternative.get('words', []):
                    total_confidence += word.get('confidence', 0)
                    confidence_count += 1
        average_confidence = (total_confidence / confidence_count) * 100
        job_data.append({"name": "Average Confidence", "value": f"{average_confidence:.2f}%"})

    for next_row in job_data:
        row_cells = table.add_row().cells
        row_cells[0].text = next_row["name"]
        row_cells[1].text = next_row["value"]

    print_transcript(response, document, args.speakers)

    if args.outputFile is None:
        args.outputFile = args.file + ".docx"
    document.save(args.outputFile)

if __name__ == "__main__":
    main()
