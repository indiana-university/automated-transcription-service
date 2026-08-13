import statistics
from datetime import timedelta
from io import BytesIO

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Mm, Pt


START_NEW_SEGMENT_SECONDS = 2.0


def _seconds(ticks):
    return float(ticks or 0) / 10_000_000


def _timestamp(ticks):
    value = timedelta(seconds=_seconds(ticks))
    total = int(value.total_seconds())
    return f"{total // 3600:02d}:{(total % 3600) // 60:02d}:{total % 60:02d}"


def transcript_summary(data, fallback_locale=None):
    phrases = [item for item in data.get("recognizedPhrases", []) if item.get("recognitionStatus") == "Success"]
    confidences = [item["nBest"][0].get("confidence", 0) for item in phrases if item.get("nBest")]
    languages = sorted({item["locale"] for item in phrases if item.get("locale")})
    if not languages and fallback_locale:
        languages = [fallback_locale]
    return {
        "confidence": round(statistics.mean(confidences) * 100, 2) if confidences else 0,
        "duration": round(_seconds(data.get("durationInTicks")), 2),
        "languages": ", ".join(languages),
        "phrases": len(phrases),
    }


def _transcript_segments(data):
    segments = []
    for phrase in data.get("recognizedPhrases", []):
        if phrase.get("recognitionStatus") != "Success" or not phrase.get("nBest"):
            continue
        speaker = phrase.get("speaker")
        if speaker is not None:
            identity = ("speaker", int(speaker))
            label = f"Speaker {int(speaker)}"
        else:
            channel = int(phrase.get("channel", 0))
            identity = ("channel", channel)
            label = f"Channel {channel + 1}"
        start = int(phrase.get("offsetInTicks") or 0)
        end = start + int(phrase.get("durationInTicks") or 0)
        if (
            segments
            and segments[-1]["identity"] == identity
            and _seconds(start - segments[-1]["end"]) < START_NEW_SEGMENT_SECONDS
        ):
            segments[-1]["end"] = max(segments[-1]["end"], end)
            segments[-1]["parts"].append(phrase["nBest"][0])
        else:
            segments.append({
                "identity": identity,
                "label": label,
                "start": start,
                "end": end,
                "parts": [phrase["nBest"][0]],
            })
    return segments


def create_docx(data, job_name, title="Transcription Results", threshold=90, fallback_locale=None):
    document = Document()
    section = document.sections[0]
    section.left_margin = section.right_margin = Mm(19.1)
    section.top_margin = section.bottom_margin = Mm(19.1)
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading(title, 1)

    summary = transcript_summary(data, fallback_locale)
    table = document.add_table(rows=0, cols=2)
    table.style = "Light List"
    for label, value in (
        ("Job Name", job_name),
        ("Audio Duration", f"{summary['duration']} seconds"),
        ("Language(s)", summary["languages"] or "Not reported"),
        ("Average Confidence", f"{summary['confidence']}%"),
    ):
        cells = table.add_row().cells
        cells[0].text, cells[1].text = label, str(value)

    document.add_heading("Transcript", 2)
    for segment in _transcript_segments(data):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"[{_timestamp(segment['start'])}] {segment['label']}: ").bold = True
        for index, best in enumerate(segment["parts"]):
            text = best.get("display") or best.get("lexical", "")
            run = paragraph.add_run((" " if index else "") + text)
            if float(best.get("confidence", 0)) * 100 < threshold:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    output = BytesIO()
    document.save(output)
    return output.getvalue(), summary
