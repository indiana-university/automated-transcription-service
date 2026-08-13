from io import BytesIO
import sys
from pathlib import Path
from zipfile import ZipFile

from docx import Document

sys.path.insert(0, str(Path(__file__).parents[1] / "src" / "functions"))

from ats.documents import create_docx, transcript_summary

SAMPLE = {
    "durationInTicks": 25_000_000,
    "recognizedPhrases": [
        {
            "recognitionStatus": "Success",
            "channel": 0,
            "speaker": 1,
            "locale": "en-US",
            "offsetInTicks": 7_600_000,
            "nBest": [{"confidence": 0.75, "display": "Hello world."}],
        },
        {
            "recognitionStatus": "Success",
            "channel": 0,
            "speaker": 2,
            "locale": "en-US",
            "offsetInTicks": 15_000_000,
            "nBest": [{"confidence": 0.95, "display": "Second speaker."}],
        },
    ],
}


def test_summary_uses_azure_ticks_and_confidence():
    assert transcript_summary(SAMPLE) == {"confidence": 85, "duration": 2.5, "languages": "en-US", "phrases": 2}


def test_summary_uses_job_locale_when_results_omit_it():
    data = {
        **SAMPLE,
        "recognizedPhrases": [
            {key: value for key, value in phrase.items() if key != "locale"}
            for phrase in SAMPLE["recognizedPhrases"]
        ],
    }

    assert transcript_summary(data, "en-US")["languages"] == "en-US"


def test_summary_prefers_identified_phrase_locales_to_fallback():
    assert transcript_summary(SAMPLE, "fr-FR")["languages"] == "en-US"


def test_docx_contains_transcript_and_speaker_label():
    content, _ = create_docx(SAMPLE, "interview", threshold=90)
    with ZipFile(__import__("io").BytesIO(content)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "Hello world." in xml
    assert "Speaker 1" in xml
    assert "Second speaker." in xml
    assert "Speaker 2" in xml
    assert "Speaker 3" not in xml
    assert "interview" in xml


def test_docx_contains_job_locale_when_results_omit_it():
    data = {
        **SAMPLE,
        "recognizedPhrases": [
            {key: value for key, value in phrase.items() if key != "locale"}
            for phrase in SAMPLE["recognizedPhrases"]
        ],
    }

    content, summary = create_docx(data, "interview", fallback_locale="en-US")

    with ZipFile(__import__("io").BytesIO(content)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "en-US" in xml
    assert "Not reported" not in xml
    assert summary["languages"] == "en-US"


def test_docx_channel_fallback_is_one_based():
    phrase = SAMPLE["recognizedPhrases"][0].copy()
    phrase.pop("speaker")
    phrase["channel"] = 0

    content, _ = create_docx({"recognizedPhrases": [phrase]}, "interview")

    with ZipFile(__import__("io").BytesIO(content)) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert "Channel 1" in xml


def test_docx_merges_contiguous_phrases_from_same_speaker():
    phrases = [
        {
            "recognitionStatus": "Success",
            "speaker": 1,
            "offsetInTicks": 2_800_000,
            "durationInTicks": 19_200_000,
            "nBest": [{"confidence": 0.95, "display": "Hello, my name is Charlie Duke."}],
        },
        {
            "recognitionStatus": "Success",
            "speaker": 1,
            "offsetInTicks": 24_800_000,
            "durationInTicks": 62_800_000,
            "nBest": [{"confidence": 0.95, "display": "I was born in Charlotte."}],
        },
    ]

    content, _ = create_docx({"recognizedPhrases": phrases}, "interview")
    transcript = [
        paragraph.text
        for paragraph in Document(BytesIO(content)).paragraphs
        if paragraph.text.startswith("[")
    ]

    assert transcript == [
        "[00:00:00] Speaker 1: Hello, my name is Charlie Duke. I was born in Charlotte."
    ]


def test_docx_starts_new_segment_after_two_second_pause():
    phrases = [
        {
            "recognitionStatus": "Success",
            "speaker": 1,
            "offsetInTicks": 0,
            "durationInTicks": 10_000_000,
            "nBest": [{"confidence": 0.95, "display": "First phrase."}],
        },
        {
            "recognitionStatus": "Success",
            "speaker": 1,
            "offsetInTicks": 30_000_000,
            "durationInTicks": 10_000_000,
            "nBest": [{"confidence": 0.95, "display": "Second phrase."}],
        },
    ]

    content, _ = create_docx({"recognizedPhrases": phrases}, "interview")
    transcript = [
        paragraph.text
        for paragraph in Document(BytesIO(content)).paragraphs
        if paragraph.text.startswith("[")
    ]

    assert transcript == [
        "[00:00:00] Speaker 1: First phrase.",
        "[00:00:03] Speaker 1: Second phrase.",
    ]
